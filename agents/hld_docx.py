"""
agents/hld_docx.py — Renders HLD content dict into a formatted .docx file.
Requires: pip install python-docx
"""

from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1a, 0x1a, 0x2e)
MID_BLUE    = RGBColor(0x16, 0x21, 0x3e)
ACCENT_BLUE = RGBColor(0x4f, 0x46, 0xe5)
LIGHT_GREY  = RGBColor(0xf5, 0xf5, 0xf7)
TEXT_GREY   = RGBColor(0x55, 0x55, 0x55)
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)


# ── Helpers ────────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def add_horizontal_rule(doc, color_hex="4f46e5", thickness=12):
    """Add a coloured horizontal rule paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), str(thickness))
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(6)
    return p


def heading(doc, text: str, level: int = 1):
    """Add a styled heading."""
    p = doc.add_heading(text, level=level)
    run = p.runs[0] if p.runs else p.add_run(text)
    if level == 1:
        run.font.color.rgb = DARK_BLUE
        run.font.size = Pt(16)
        run.bold = True
        add_horizontal_rule(doc)
    elif level == 2:
        run.font.color.rgb = ACCENT_BLUE
        run.font.size = Pt(13)
        run.bold = True
    return p


def body_text(doc, text: str, italic: bool = False, color: RGBColor = None):
    """Add a styled body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    if italic:
        run.italic = True
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(6)
    return p


def bullet_item(doc, text: str, bold_prefix: str = None):
    """Add a bullet list item, with optional bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(f"{bold_prefix} ")
        run.bold = True
        run.font.size = Pt(11)
        run.font.name = 'Arial'
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(3)
    return p


def add_cover_page(doc, title: str, version: str, status: str, date: str):
    """Add a cover page with title block."""
    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.font.size = Pt(26)
    run.font.bold = True
    run.font.color.rgb = DARK_BLUE
    run.font.name = 'Arial'

    doc.add_paragraph()

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("High Level Design")
    run.font.size = Pt(16)
    run.font.color.rgb = ACCENT_BLUE
    run.font.name = 'Arial'

    for _ in range(2):
        doc.add_paragraph()

    # Document control table
    table = doc.add_table(rows=4, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    rows_data = [
        ("Version", version),
        ("Status",  status),
        ("Date",    date),
        ("Classification", "Internal"),
    ]

    for i, (label, value) in enumerate(rows_data):
        row = table.rows[i]
        # Label cell
        label_cell = row.cells[0]
        set_cell_bg(label_cell, "1a1a2e")
        lp = label_cell.paragraphs[0]
        lr = lp.add_run(label)
        lr.font.bold = True
        lr.font.color.rgb = WHITE
        lr.font.size = Pt(10)
        lr.font.name = 'Arial'
        label_cell.width = Inches(2)

        # Value cell
        value_cell = row.cells[1]
        set_cell_bg(value_cell, "f5f5f7")
        vp = value_cell.paragraphs[0]
        vr = vp.add_run(value)
        vr.font.size = Pt(10)
        vr.font.name = 'Arial'
        value_cell.width = Inches(3)

    doc.add_page_break()


def add_risks_table(doc, risks: list):
    """Add a formatted risks and open questions table."""
    if not risks:
        body_text(doc, "No risks or open questions identified.", italic=True, color=TEXT_GREY)
        return

    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_cells = table.rows[0].cells
    headers = ["Type", "Description", "Owner"]
    col_widths = [Inches(1.2), Inches(4.0), Inches(1.5)]

    for i, (cell, header) in enumerate(zip(hdr_cells, headers)):
        set_cell_bg(cell, "1a1a2e")
        cell.width = col_widths[i]
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(10)
        run.font.name = 'Arial'

    for item in risks:
        row = table.add_row()
        row.cells[0].width = col_widths[0]
        row.cells[1].width = col_widths[1]
        row.cells[2].width = col_widths[2]

        item_type = item.get("type", "")
        bg = "fff3cd" if "question" in item_type.lower() else "f8d7da"
        set_cell_bg(row.cells[0], bg)

        for cell, key in zip(row.cells, ["type", "description", "owner"]):
            p = cell.paragraphs[0]
            run = p.add_run(item.get(key, "TBC"))
            run.font.size = Pt(10)
            run.font.name = 'Arial'

    doc.add_paragraph()


def render_hld_docx(hld: dict, output_path: str) -> str:
    """
    Render an HLD content dict into a formatted .docx file.
    Returns the output path.
    """
    doc = Document()

    # ── Page setup ─────────────────────────────────────────────────────────────
    section = doc.sections[0]
    section.page_width  = Inches(8.27)   # A4
    section.page_height = Inches(11.69)
    section.left_margin   = Inches(1.0)
    section.right_margin  = Inches(1.0)
    section.top_margin    = Inches(1.0)
    section.bottom_margin = Inches(1.0)

    # ── Default font ───────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)

    # ── Cover page ─────────────────────────────────────────────────────────────
    dc = hld.get("document_control", {})
    add_cover_page(
        doc,
        title=dc.get("title", "High Level Design"),
        version=dc.get("version", "0.1"),
        status=dc.get("status", "Draft"),
        date=dc.get("date", datetime.now().strftime("%d %B %Y")),
    )

    # ── 1. Executive Summary ───────────────────────────────────────────────────
    heading(doc, "1. Executive Summary", level=1)
    body_text(doc, hld.get("executive_summary", ""))

    # ── 2. Background ──────────────────────────────────────────────────────────
    heading(doc, "2. Background & Problem Statement", level=1)
    body_text(doc, hld.get("background", ""))

    # ── 3. Current State ───────────────────────────────────────────────────────
    heading(doc, "3. Current State", level=1)
    cs = hld.get("current_state", {})
    body_text(doc, cs.get("description", ""))
    limitations = cs.get("limitations", [])
    if limitations:
        heading(doc, "Limitations", level=2)
        for lim in limitations:
            bullet_item(doc, lim)

    # ── 4. Proposed Architecture ───────────────────────────────────────────────
    heading(doc, "4. Proposed Architecture", level=1)
    pa = hld.get("proposed_architecture", {})
    body_text(doc, pa.get("description", ""))

    components = pa.get("components", [])
    if components:
        heading(doc, "Components", level=2)
        for comp in components:
            bullet_item(doc, comp.get("description", ""), bold_prefix=comp.get("name", ""))

    key_flows = pa.get("key_flows", [])
    if key_flows:
        heading(doc, "Key Flows", level=2)
        for flow in key_flows:
            bullet_item(doc, flow)

    # Note about diagram
    p = doc.add_paragraph()
    run = p.add_run("ℹ️  The full architecture diagram is available in the Meeting Artefact Pipeline UI (Architecture Diagram tab) and as a downloadable .mmd file.")
    run.font.italic = True
    run.font.color.rgb = TEXT_GREY
    run.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(6)

    # ── 5. Design Decisions ────────────────────────────────────────────────────
    heading(doc, "5. Key Design Decisions", level=1)
    decisions = hld.get("design_decisions", [])
    if decisions:
        for dec in decisions:
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(dec.get("decision", ""))
            run.font.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Arial'
            rationale = dec.get("rationale", "")
            if rationale:
                p2 = doc.add_paragraph()
                p2.paragraph_format.left_indent = Inches(0.4)
                r2 = p2.add_run(f"Rationale: {rationale}")
                r2.font.italic = True
                r2.font.size = Pt(10)
                r2.font.color.rgb = TEXT_GREY
                r2.font.name = 'Arial'
                p2.paragraph_format.space_after = Pt(4)
    else:
        body_text(doc, "No decisions were recorded for this meeting.", italic=True, color=TEXT_GREY)

    # ── 6. Risks & Open Questions ──────────────────────────────────────────────
    heading(doc, "6. Risks & Open Questions", level=1)
    add_risks_table(doc, hld.get("risks_and_open_questions", []))

    # ── 7. Next Steps ──────────────────────────────────────────────────────────
    heading(doc, "7. Next Steps & Actions", level=1)
    next_steps = hld.get("next_steps", [])
    if next_steps:
        for step in next_steps:
            bullet_item(doc, step.get("action", ""), bold_prefix=f"[{step.get('owner', 'UNASSIGNED')}]")
    else:
        body_text(doc, "No actions recorded.", italic=True, color=TEXT_GREY)

    # ── Footer ─────────────────────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run(f"{dc.get('title', 'HLD')}  |  {dc.get('version', '0.1')}  |  {dc.get('status', 'Draft')}  |  Internal")
        run.font.size = Pt(9)
        run.font.color.rgb = TEXT_GREY
        run.font.name = 'Arial'

    doc.save(output_path)
    return output_path